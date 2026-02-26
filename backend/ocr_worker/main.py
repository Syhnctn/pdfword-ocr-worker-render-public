import io
import os
import re
from datetime import datetime, timezone
from typing import Any

import httpx
from docx import Document
from fastapi import FastAPI, HTTPException, Request
from pydantic import BaseModel
from pypdf import PdfReader
from supabase import Client, create_client

try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover - optional dependency
    fitz = None

try:
    import pytesseract
    from pytesseract import Output
    from PIL import Image
    from PIL import ImageFilter, ImageOps
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None
    Output = None
    Image = None
    ImageFilter = None
    ImageOps = None

app = FastAPI(title="pdfword-ocr-worker", version="0.2.0")

_BULLET_RE = re.compile(r"^([\-*]|\d+[.)])\s+")
_MULTISPACE_RE = re.compile(r"\s+")
_PUNCT_END_RE = re.compile(r"[.!?:;)](?:['\"])?$")
_HEADING_RE = re.compile(r"^[A-Z0-9][A-Z0-9\s/&()_-]{2,}$")
_WORD_RE = re.compile(r"\w+", re.UNICODE)
_TURKISH_CHARS = "çğıöşüÇĞİÖŞÜ"


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def env_flag(name: str, default: bool) -> bool:
    raw = os.environ.get(name, "").strip().lower()
    if not raw:
        return default
    return raw not in {"0", "false", "no", "off"}


def open_source_ocr_enabled() -> bool:
    return env_flag("OPEN_SOURCE_OCR_ENABLED", True)


def tesseract_langs() -> str:
    value = os.environ.get("TESSERACT_LANG", "tur+eng").strip()
    return value or "tur+eng"


def tesseract_dpi() -> int:
    raw = os.environ.get("TESSERACT_DPI", "220").strip()
    try:
        dpi = int(raw)
    except ValueError:
        dpi = 220
    return max(96, min(dpi, 600))


def tesseract_psm() -> str:
    raw = os.environ.get("TESSERACT_PSM", "6").strip()
    return raw or "6"


def tesseract_oem() -> str:
    raw = os.environ.get("TESSERACT_OEM", "1").strip()
    return raw if raw in {"0", "1", "2", "3"} else "1"


def tesseract_psm_candidates() -> list[str]:
    raw = os.environ.get("TESSERACT_PSM_CANDIDATES", "").strip()
    values = raw.split(",") if raw else [tesseract_psm(), "4"]
    seen: set[str] = set()
    result: list[str] = []
    for part in values:
        item = part.strip()
        if not item or not item.isdigit() or item in seen:
            continue
        seen.add(item)
        result.append(item)
    return result or ["6", "4"]


def tesseract_max_variants() -> int:
    raw = os.environ.get("TESSERACT_MAX_VARIANTS", "3").strip()
    try:
        value = int(raw)
    except ValueError:
        value = 3
    return max(1, min(value, 12))


def tesseract_call_timeout_sec() -> float:
    raw = os.environ.get("TESSERACT_CALL_TIMEOUT_SEC", "8").strip()
    try:
        value = float(raw)
    except ValueError:
        value = 8.0
    return max(2.0, min(value, 60.0))


def tesseract_max_attempts() -> int:
    raw = os.environ.get("TESSERACT_MAX_ATTEMPTS", "4").strip()
    try:
        value = int(raw)
    except ValueError:
        value = 4
    return max(1, min(value, 20))


def make_supabase_client() -> Client:
    url = os.environ.get("SUPABASE_URL", "").strip()
    key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "").strip()
    if not url or not key:
        raise RuntimeError("SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY are required")
    return create_client(url, key)


def assert_worker_secret(request: Request) -> None:
    expected = os.environ.get("OCR_WORKER_SECRET", "").strip()
    if not expected:
        return

    auth = request.headers.get("Authorization", "")
    token = auth.replace("Bearer ", "").strip()
    if not token or token != expected:
        raise HTTPException(status_code=401, detail="invalid_worker_secret")


def sanitize_storage_name(name: str) -> str:
    safe = re.sub(r"[^a-zA-Z0-9._-]", "_", name.strip())
    return safe or "input.pdf"


def coerce_input_files(
    job_id: str, input_meta: list[dict[str, Any]]
) -> list[dict[str, Any]]:
    input_bucket = os.environ.get("OCR_INPUT_BUCKET", "ocr-inputs")
    files: list[dict[str, Any]] = []

    for index, item in enumerate(input_meta):
        if not isinstance(item, dict):
            continue

        name = str(item.get("name", f"input-{index}.pdf"))
        path = str(item.get("path") or item.get("storage_path") or "").strip()
        bucket = str(item.get("bucket") or item.get("storage_bucket") or input_bucket)
        mime_type = str(item.get("mime_type") or "application/pdf")

        if not path:
            legacy_name = sanitize_storage_name(name)
            path = f"{job_id}/input/{index:02d}_{legacy_name}"

        files.append(
            {
                "name": name,
                "bucket": bucket,
                "path": path,
                "mime_type": mime_type,
                "size_mb": item.get("size_mb"),
            }
        )

    return files


def build_placeholder_markdown(input_meta: list[dict[str, Any]]) -> str:
    lines = ["# OCR Result", ""]
    lines.append("This output was generated by backend worker.")
    lines.append("")
    lines.append("## Input Files")
    for item in input_meta:
        name = str(item.get("name", "unknown.pdf"))
        size_mb = item.get("size_mb", 0)
        lines.append(f"- {name} ({size_mb} MB)")
    lines.append("")
    lines.append("## Extracted Text")
    lines.append(
        "No readable text could be extracted locally. For scanned/image PDFs, enable open-source OCR (Tesseract) or configure LIGHTON_OCR_ENDPOINT."
    )
    return "\n".join(lines)


def normalize_extracted_text(raw_text: str) -> str:
    normalized = raw_text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.strip() for line in normalized.split("\n")]
    paragraphs: list[str] = []
    current = ""

    def flush() -> None:
        nonlocal current
        if current:
            paragraphs.append(current.strip())
            current = ""

    for line in lines:
        if not line:
            flush()
            continue

        line = _MULTISPACE_RE.sub(" ", line)

        if not current:
            current = line
            continue

        if current.endswith("-") and line and line[0].islower():
            current = current[:-1] + line
            continue

        current_is_list = bool(_BULLET_RE.match(current))
        next_is_list = bool(_BULLET_RE.match(line))
        current_is_heading = len(current) <= 80 and bool(_HEADING_RE.match(current))
        next_is_heading = len(line) <= 80 and bool(_HEADING_RE.match(line))
        current_ends_sentence = bool(_PUNCT_END_RE.search(current))

        if current_is_list or next_is_list or current_is_heading or next_is_heading:
            flush()
            current = line
            continue

        if current_ends_sentence:
            flush()
            current = line
            continue

        current = f"{current} {line}"

    flush()

    deduped: list[str] = []
    prev = None
    for paragraph in paragraphs:
        if not paragraph:
            continue
        if paragraph == prev:
            continue
        deduped.append(paragraph)
        prev = paragraph

    return "\n\n".join(deduped)


def extract_pdf_text_sections(pdf_bytes: bytes) -> list[tuple[int, str]]:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # pragma: no cover - depends on file
            raise RuntimeError(f"encrypted_pdf:{exc}") from exc

    sections: list[tuple[int, str]] = []
    for page_index, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        text = normalize_extracted_text(raw)
        if text:
            sections.append((page_index, text))

    return sections


def _otsu_threshold(gray_image: Any) -> int:
    histogram = gray_image.histogram()
    if not histogram or len(histogram) < 256:
        return 180

    counts = histogram[:256]
    total = sum(counts)
    if total <= 0:
        return 180

    weighted_sum = sum(index * count for index, count in enumerate(counts))
    sum_b = 0.0
    weight_b = 0
    best_variance = -1.0
    threshold = 180

    for index, count in enumerate(counts):
        weight_b += count
        if weight_b == 0:
            continue
        weight_f = total - weight_b
        if weight_f == 0:
            break

        sum_b += index * count
        mean_b = sum_b / weight_b
        mean_f = (weighted_sum - sum_b) / weight_f
        variance = weight_b * weight_f * ((mean_b - mean_f) ** 2)

        if variance > best_variance:
            best_variance = variance
            threshold = index

    return max(40, min(threshold, 230))


def _binarize_luma(gray_image: Any) -> Any:
    threshold = _otsu_threshold(gray_image)
    return gray_image.point(lambda px, t=threshold: 255 if px >= t else 0, mode="L")


def _lanczos_resample() -> int:
    if Image is None:
        return 1
    resampling = getattr(Image, "Resampling", None)
    if resampling is not None:
        return int(resampling.LANCZOS)
    return int(getattr(Image, "LANCZOS", 1))


def _build_tesseract_image_variants(image: Any) -> list[tuple[str, Any]]:
    if Image is None:
        return [("raw", image)]

    gray = image.convert("L")
    variants: list[tuple[str, Any]] = []

    auto = ImageOps.autocontrast(gray) if ImageOps is not None else gray
    variants.append(("gray_auto", auto))
    variants.append(("binary_auto", _binarize_luma(auto)))
    variants.append(("gray", gray))

    if ImageFilter is not None:
        denoised = auto.filter(ImageFilter.MedianFilter(size=3))
        variants.append(("gray_auto_median", denoised))
        variants.append(("binary_auto_median", _binarize_luma(denoised)))

    width, height = auto.size
    if max(width, height) < 2600:
        upscaled = auto.resize((width * 2, height * 2), _lanczos_resample())
        variants.append(("gray_auto_2x", upscaled))
        variants.append(("binary_auto_2x", _binarize_luma(upscaled)))

    deduped: list[tuple[str, Any]] = []
    seen = set()
    for label, variant in variants:
        key = (label, getattr(variant, "mode", ""), getattr(variant, "size", None))
        if key in seen:
            continue
        seen.add(key)
        deduped.append((label, variant))
    return deduped[: tesseract_max_variants()]


def _ocr_candidate_score(text: str, mean_confidence: float) -> float:
    normalized = normalize_extracted_text(text)
    if not normalized:
        return -1e9

    text_len = len(normalized)
    word_count = len(_WORD_RE.findall(normalized))
    turkish_hits = sum(normalized.count(ch) for ch in _TURKISH_CHARS)
    replacement_hits = normalized.count("\ufffd")
    symbol_noise = normalized.count("|") + normalized.count("~")

    return (
        (mean_confidence * 4.0)
        + float(text_len)
        + float(word_count * 2)
        + float(turkish_hits * 8)
        - float(replacement_hits * 25)
        - float(symbol_noise * 4)
    )


def _ocr_candidate_good_enough(text: str, mean_confidence: float) -> bool:
    normalized = normalize_extracted_text(text)
    if len(normalized) < 48:
        return False
    if mean_confidence >= 75:
        return True
    turkish_hits = sum(normalized.count(ch) for ch in _TURKISH_CHARS)
    return turkish_hits >= 3 and mean_confidence >= 50


def _tesseract_config(psm: str) -> str:
    return (
        f"--oem {tesseract_oem()} --psm {psm} "
        "-c preserve_interword_spaces=1"
    )


def _run_tesseract_candidate(image: Any, lang: str, psm: str) -> tuple[str, float]:
    if pytesseract is None:
        return "", -1.0

    config = _tesseract_config(psm)
    timeout_sec = tesseract_call_timeout_sec()

    text = ""
    mean_conf = -1.0

    if Output is not None:
        try:
            data = pytesseract.image_to_data(
                image,
                lang=lang,
                config=config,
                output_type=Output.DICT,
                timeout=timeout_sec,
            )
            tokens: list[str] = []
            confs: list[float] = []
            for token, conf in zip(data.get("text", []), data.get("conf", [])):
                token_str = str(token or "").strip()
                if token_str:
                    tokens.append(token_str)
                try:
                    conf_val = float(conf)
                except (TypeError, ValueError):
                    continue
                if conf_val >= 0:
                    confs.append(conf_val)
            text = " ".join(tokens)
            if confs:
                mean_conf = sum(confs) / len(confs)
        except Exception:
            text = ""
            mean_conf = -1.0

    if not text:
        try:
            text = pytesseract.image_to_string(
                image, lang=lang, config=config, timeout=timeout_sec
            )
        except TypeError:
            text = pytesseract.image_to_string(image, lang=lang, config=config)
        except Exception:
            text = ""

    return text, mean_conf


def extract_pdf_text_sections_with_tesseract(pdf_bytes: bytes) -> list[tuple[int, str]]:
    if fitz is None or pytesseract is None or Image is None:
        raise RuntimeError("open_source_ocr_dependencies_missing")

    dpi = tesseract_dpi()
    zoom = dpi / 72.0
    lang = tesseract_langs()
    psm_candidates = tesseract_psm_candidates()
    sections: list[tuple[int, str]] = []

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:  # pragma: no cover - dependency/file specific
        raise RuntimeError(f"open_source_ocr_pdf_open_failed:{exc}") from exc

    try:
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png")))
            try:
                best_text = ""
                best_score = -1e9
                best_variant: Any | None = None
                best_psm = psm_candidates[0] if psm_candidates else "6"

                variants = _build_tesseract_image_variants(image)
                attempts: list[tuple[Any, str]] = []
                if psm_candidates:
                    primary_psm = psm_candidates[0]
                    for _, variant in variants:
                        attempts.append((variant, primary_psm))
                    for extra_psm in psm_candidates[1:]:
                        for _, variant in variants[:2]:
                            attempts.append((variant, extra_psm))

                attempts = attempts[: tesseract_max_attempts()]

                for variant, psm in attempts:
                    candidate_text, candidate_conf = _run_tesseract_candidate(
                        variant, lang, psm
                    )
                    score = _ocr_candidate_score(candidate_text, candidate_conf)
                    if score > best_score:
                        best_score = score
                        best_text = candidate_text
                        best_variant = variant
                        best_psm = psm
                    if _ocr_candidate_good_enough(candidate_text, candidate_conf):
                        break

                if not best_text.strip():
                    raise RuntimeError("tesseract_ocr_empty_result")
                raw = best_text
                if best_variant is not None:
                    try:
                        pretty = pytesseract.image_to_string(
                            best_variant,
                            lang=lang,
                            config=_tesseract_config(best_psm),
                            timeout=tesseract_call_timeout_sec(),
                        )
                        if pretty.strip():
                            raw = pretty
                    except Exception:
                        pass
            except Exception as exc:  # pragma: no cover - tesseract specific
                raise RuntimeError(f"tesseract_ocr_failed:{exc}") from exc
            finally:
                image.close()

            text = normalize_extracted_text(raw or "")
            if text:
                sections.append((page_index + 1, text))
    finally:
        doc.close()

    return sections


def download_storage_bytes(sb: Client, bucket: str, path: str) -> bytes:
    result = sb.storage.from_(bucket).download(path)

    if isinstance(result, (bytes, bytearray)):
        return bytes(result)

    if isinstance(result, tuple) and result:
        first = result[0]
        if isinstance(first, (bytes, bytearray)):
            return bytes(first)

    if hasattr(result, "content"):
        return bytes(result.content)

    raise RuntimeError(f"download_failed_unexpected_type:{type(result).__name__}")


def build_markdown_from_extracted_files(file_results: list[dict[str, Any]]) -> str:
    lines = ["# OCR Result", ""]
    lines.append("This output was generated by backend worker.")
    lines.append("")

    for item in file_results:
        name = str(item.get("name", "unknown.pdf"))
        lines.append(f"## {name}")
        lines.append("")

        note = str(item.get("note") or "").strip()
        if note:
            lines.append(note)
            lines.append("")

        sections = item.get("sections") or []
        if isinstance(sections, list) and sections:
            multi_page = len(sections) > 1
            for page_num, text in sections:
                if multi_page:
                    lines.append(f"### Page {page_num}")
                    lines.append("")
                lines.append(str(text))
                lines.append("")
            continue

        lines.append(
            "No readable embedded text was found in this PDF. It may be scanned/image-based and needs OCR (Tesseract or external OCR)."
        )
        lines.append("")

    return "\n".join(lines).strip() + "\n"


async def call_lighton_ocr_endpoint(
    endpoint_url: str, token: str, input_meta: list[dict[str, Any]]
) -> str:
    payload = {
        "inputs": {
            "prompt": "Extract text from provided document pages.",
            "files": input_meta,
        }
    }
    async with httpx.AsyncClient(timeout=120) as client:
        response = await client.post(
            endpoint_url,
            json=payload,
            headers={"Authorization": f"Bearer {token}"},
        )
        response.raise_for_status()
        data = response.json()
        if isinstance(data, dict):
            if isinstance(data.get("markdown"), str):
                return data["markdown"]
            if isinstance(data.get("text"), str):
                return data["text"]
        return str(data)


def markdown_to_docx_bytes(markdown: str) -> bytes:
    document = Document()
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def upload_outputs(
    sb: Client, job_id: str, markdown: str, docx_bytes: bytes
) -> tuple[str, str]:
    bucket = os.environ.get("OCR_RESULTS_BUCKET", "ocr-results")
    md_path = f"{job_id}/result.md"
    docx_path = f"{job_id}/result.docx"

    sb.storage.from_(bucket).upload(
        md_path,
        markdown.encode("utf-8"),
        file_options={"content-type": "text/markdown", "upsert": "true"},
    )
    sb.storage.from_(bucket).upload(
        docx_path,
        docx_bytes,
        file_options={
            "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "upsert": "true",
        },
    )
    return md_path, docx_path


async def process_job(job_id: str) -> dict[str, Any]:
    sb = make_supabase_client()
    row_res = (
        sb.table("ocr_jobs")
        .select("id, input_meta, status")
        .eq("id", job_id)
        .limit(1)
        .execute()
    )
    rows = row_res.data or []
    if not rows:
        raise HTTPException(status_code=404, detail="job_not_found")

    input_meta = rows[0].get("input_meta") or []
    if not isinstance(input_meta, list):
        input_meta = []

    sb.table("ocr_jobs").update(
        {"status": "processing", "progress_pct": 25, "started_at": utc_now()}
    ).eq("id", job_id).execute()

    try:
        resolved_files = coerce_input_files(job_id, input_meta)
        extracted_files: list[dict[str, Any]] = []
        has_local_text = False
        use_oss_ocr = open_source_ocr_enabled()

        for index, file_meta in enumerate(resolved_files):
            name = str(file_meta.get("name", f"input-{index}.pdf"))
            bucket = str(file_meta.get("bucket", "ocr-inputs"))
            path = str(file_meta.get("path", "")).strip()

            progress = 35 + int(((index + 1) / max(len(resolved_files), 1)) * 25)
            sb.table("ocr_jobs").update({"progress_pct": progress}).eq(
                "id", job_id
            ).execute()

            if not path:
                extracted_files.append(
                    {
                        "name": name,
                        "sections": [],
                        "note": "Storage path is missing for this file.",
                    }
                )
                continue

            try:
                pdf_bytes = download_storage_bytes(sb, bucket, path)
                sections = extract_pdf_text_sections(pdf_bytes)
                note = ""

                if not sections and use_oss_ocr:
                    try:
                        sections = extract_pdf_text_sections_with_tesseract(pdf_bytes)
                        if sections:
                            note = "Extracted with open-source OCR (Tesseract)."
                    except Exception as exc:  # pragma: no cover - env dependent
                        note = f"Open-source OCR unavailable: {exc}"

                has_local_text = has_local_text or bool(sections)
                item: dict[str, Any] = {"name": name, "sections": sections}
                if note:
                    item["note"] = note
                extracted_files.append(item)
            except Exception as exc:  # pragma: no cover - storage/pdf dependent
                extracted_files.append(
                    {
                        "name": name,
                        "sections": [],
                        "note": f"Could not read PDF content: {exc}",
                    }
                )

        endpoint_url = os.environ.get("LIGHTON_OCR_ENDPOINT", "").strip()
        endpoint_token = os.environ.get("LIGHTON_OCR_TOKEN", "").strip()

        if has_local_text:
            markdown = build_markdown_from_extracted_files(extracted_files)
        elif endpoint_url and endpoint_token:
            markdown = await call_lighton_ocr_endpoint(
                endpoint_url, endpoint_token, resolved_files or input_meta
            )
        else:
            markdown = build_markdown_from_extracted_files(extracted_files)
            if not markdown.strip():
                markdown = build_placeholder_markdown(input_meta)

        sb.table("ocr_jobs").update({"progress_pct": 80}).eq("id", job_id).execute()

        docx_bytes = markdown_to_docx_bytes(markdown)
        md_path, docx_path = upload_outputs(sb, job_id, markdown, docx_bytes)

        sb.table("ocr_jobs").update(
            {
                "status": "succeeded",
                "progress_pct": 100,
                "output_md_path": md_path,
                "output_docx_path": docx_path,
                "finished_at": utc_now(),
            }
        ).eq("id", job_id).execute()

        return {
            "job_id": job_id,
            "status": "succeeded",
            "output_md_path": md_path,
            "output_docx_path": docx_path,
        }
    except Exception as exc:
        sb.table("ocr_jobs").update(
            {
                "status": "failed",
                "progress_pct": 100,
                "error_code": "ocr_error",
                "error_message": str(exc),
                "finished_at": utc_now(),
            }
        ).eq("id", job_id).execute()
        raise


@app.get("/healthz")
def healthz() -> dict[str, str]:
    return {"status": "ok"}


class ProcessRequest(BaseModel):
    job_id: str


@app.post("/internal/process")
async def process(request: Request, payload: ProcessRequest) -> dict[str, Any]:
    assert_worker_secret(request)
    return await process_job(payload.job_id)


@app.post("/internal/process/{job_id}")
async def process_with_path(job_id: str, request: Request) -> dict[str, Any]:
    assert_worker_secret(request)
    return await process_job(job_id)
